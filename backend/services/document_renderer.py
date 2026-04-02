"""Phase 7 Document Renderer — produces branded .docx/.pptx from markdown + templates.

New approach (v2):
  - DOCX: Opens template, extracts cover page (up to in-body sectPr), replaces
    client/date placeholders on cover and footer, deletes body placeholder stubs,
    tokenizes the Phase 5 .md file, converts markdown blocks to python-docx
    elements using the template's built-in style IDs, inserts chart PNGs from
    Supabase ``engagement-graphics`` bucket, and appends content between the
    cover page and final sectPr.
  - PPTX: Opens template, keeps title slide, creates new slides from markdown
    sections with chart images and speaker notes.
"""
from __future__ import annotations

import io
import logging
import os
import re
import zipfile
from copy import deepcopy
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from lxml import etree

from services.supabase_client import get_supabase, get_engagement_by_id, log_activity
from services.google_drive_engagement import (
    list_files_in_folder,
    download_file_by_id,
    upload_file_to_drive_folder,
)

logger = logging.getLogger("baxterlabs.document_renderer")

TEMPLATES_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "templates")

# Brand constants for table styling (headings/body use template styles)
CRIMSON_HEX = "66151C"
CHARCOAL = RGBColor(0x33, 0x33, 0x33)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BODY_FONT = "Calibri"

# Map file_prefix to template filenames, output format, and output_number.
# output_number matches Phase 5 seed numbering for chart placement filtering.
_TEMPLATE_LIST: List[Tuple[str, str, str, str, int]] = [
    # (file_prefix, display_name, template_filename, output_extension, output_number)
    ("Executive_Summary", "Executive Summary", "10_Executive_Summary.docx", "docx", 1),
    ("Full_Diagnostic_Report", "Full Diagnostic Report", "09_Full_Diagnostic_Report.docx", "docx", 2),
    ("Presentation_Deck", "Presentation Deck", "11_Presentation_Deck.pptx", "pptx", 3),
    ("Implementation_Roadmap", "Implementation Roadmap", "26_90_Day_Implementation_Roadmap.docx", "docx", 4),
    ("Phase_2_Retainer_Proposal", "Phase 2 Retainer Proposal", "17_Phase2_Retainer_Proposal.docx", "docx", 5),
    # Alternate prefix — Cowork saves as Retainer_Proposal_* not Phase_2_Retainer_Proposal_*
    ("Retainer_Proposal", "Phase 2 Retainer Proposal", "17_Phase2_Retainer_Proposal.docx", "docx", 5),
]
_TEMPLATE_LIST.sort(key=lambda t: len(t[0]), reverse=True)


def _match_template(filename: str) -> Optional[Tuple[str, str, str, int, str]]:
    """Match a Drive filename to a template.

    Returns ``(display_name, template_filename, output_ext, output_number, file_prefix)`` or ``None``.
    """
    base = filename.rsplit(".", 1)[0] if "." in filename else filename
    for prefix, display_name, tpl, ext, out_num in _TEMPLATE_LIST:
        if base.startswith(prefix):
            return (display_name, tpl, ext, out_num, prefix)
    return None


# ---------------------------------------------------------------------------
# Markdown tokenizer (from POC, adapted for template-style rendering)
# ---------------------------------------------------------------------------

class _InlineSpan:
    """An inline text span with formatting flags."""
    __slots__ = ("text", "bold", "italic")

    def __init__(self, text: str, bold: bool = False, italic: bool = False):
        self.text = text
        self.bold = bold
        self.italic = italic


class _Token:
    """A parsed markdown block element."""
    __slots__ = ("type", "level", "children", "text", "rows", "chart_name")

    def __init__(
        self,
        type: str,
        level: int = 0,
        children: Optional[List[_InlineSpan]] = None,
        text: str = "",
        rows: Optional[List[List[str]]] = None,
        chart_name: str = "",
    ):
        self.type = type
        self.level = level
        self.children = children
        self.text = text
        self.rows = rows
        self.chart_name = chart_name


def _parse_inline(text: str) -> List[_InlineSpan]:
    """Parse **bold** and *italic* markers into spans."""
    spans: List[_InlineSpan] = []
    pattern = re.compile(
        r"(\*\*\*(.+?)\*\*\*)"
        r"|(\*\*(.+?)\*\*)"
        r"|(\*(.+?)\*)"
        r"|([^*]+)"
    )
    for m in pattern.finditer(text):
        if m.group(2):
            spans.append(_InlineSpan(m.group(2), bold=True, italic=True))
        elif m.group(4):
            spans.append(_InlineSpan(m.group(4), bold=True))
        elif m.group(6):
            spans.append(_InlineSpan(m.group(6), italic=True))
        elif m.group(7):
            spans.append(_InlineSpan(m.group(7)))
    return spans


def _parse_table_lines(lines: List[str]) -> List[List[str]]:
    """Parse pipe-delimited table lines into rows of cells."""
    rows: List[List[str]] = []
    for line in lines:
        cells = [c.strip() for c in line.split("|")]
        if cells and cells[0] == "":
            cells = cells[1:]
        if cells and cells[-1] == "":
            cells = cells[:-1]
        if all(re.match(r"^:?-+:?$", c) for c in cells):
            continue
        rows.append(cells)
    return rows


def _strip_title_frontmatter(md: str) -> str:
    """Strip the title/front-matter block from the start of markdown.

    The title block consists of lines of text (document title, client name,
    byline, date) followed by a ``---`` / ``***`` / ``___`` horizontal rule.
    This duplicates the cover page, so we strip everything through the last
    front-matter separator near the top.

    Also strips YAML frontmatter (``---`` delimited blocks at the very start).
    """
    content = md.strip()

    # Strip YAML frontmatter first (starts and ends with ---)
    if content.startswith("---"):
        end = content.find("---", 3)
        if end != -1:
            content = content[end + 3:].strip()

    # Now strip the title block: lines of text followed by a horizontal rule.
    # If there are two rules close together at the top, strip through the second.
    lines = content.split("\n")
    last_rule_idx = -1
    for i, line in enumerate(lines):
        stripped = line.strip()
        # Stop searching after the first heading or once we've gone past ~30 lines
        if i > 30:
            break
        if re.match(r"^#{1,6}\s+", stripped):
            break
        if stripped in ("---", "***", "___"):
            last_rule_idx = i

    if last_rule_idx >= 0:
        content = "\n".join(lines[last_rule_idx + 1:]).strip()

    return content


def _tokenize_markdown(md: str) -> List[_Token]:
    """Tokenize markdown into block-level tokens.

    Strips title front-matter before parsing.
    """
    content = _strip_title_frontmatter(md)

    tokens: List[_Token] = []
    lines = content.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Horizontal rule — skip entirely (don't render as text)
        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # Heading
        heading_match = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            heading_text = heading_match.group(2).strip()
            tokens.append(_Token(
                type="heading", level=level, text=heading_text,
                children=_parse_inline(heading_text),
            ))
            i += 1
            continue

        # Chart placeholder
        chart_match = re.match(r"^\[CHART:\s*([^\]]+)\]$", stripped)
        if chart_match:
            tokens.append(_Token(type="chart_placeholder", chart_name=chart_match.group(1).strip()))
            i += 1
            continue

        # Table
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            rows = _parse_table_lines(table_lines)
            if rows:
                tokens.append(_Token(type="table", rows=rows))
            continue

        # Bullet list item
        bullet_match = re.match(r"^[-*+]\s+(.+)$", stripped)
        if bullet_match:
            tokens.append(_Token(
                type="bullet", text=bullet_match.group(1).strip(),
                children=_parse_inline(bullet_match.group(1).strip()),
            ))
            i += 1
            continue

        # Numbered list item
        numbered_match = re.match(r"^\d+[.)]\s+(.+)$", stripped)
        if numbered_match:
            tokens.append(_Token(
                type="numbered", text=numbered_match.group(1).strip(),
                children=_parse_inline(numbered_match.group(1).strip()),
            ))
            i += 1
            continue

        # Paragraph (collect contiguous non-empty, non-special lines)
        para_lines = []
        while i < len(lines):
            l = lines[i].strip()
            if not l:
                break
            if re.match(r"^#{1,6}\s+", l):
                break
            if re.match(r"^\[CHART:", l):
                break
            if l.startswith("|"):
                break
            if re.match(r"^[-*+]\s+", l):
                break
            if re.match(r"^\d+[.)]\s+", l):
                break
            para_lines.append(l)
            i += 1

        if para_lines:
            para_text = " ".join(para_lines)
            tokens.append(_Token(
                type="paragraph", text=para_text,
                children=_parse_inline(para_text),
            ))

    return tokens


# ---------------------------------------------------------------------------
# Endnote processing helpers
# ---------------------------------------------------------------------------

# Unicode superscript digit mapping
_SUPERSCRIPT_DIGITS = str.maketrans("0123456789", "⁰¹²³⁴⁵⁶⁷⁸⁹")


def _to_superscript(n: int) -> str:
    """Convert an integer to Unicode superscript characters."""
    return str(n).translate(_SUPERSCRIPT_DIGITS)


def _number_endnotes(tokens: List[_Token]) -> List[_Token]:
    """Add sequential numbers to endnote entries in the Endnotes section.

    Detects the Endnotes heading, then prepends '1. ', '2. ', etc. to each
    subsequent paragraph that starts with '[Verified:', '[Estimated:',
    '[Stated:', or '[Derived:'.
    """
    in_endnotes = False
    endnote_num = 0
    citation_pattern = re.compile(r"^\[(Verified|Estimated|Stated|Derived):")

    for token in tokens:
        if token.type == "heading" and token.text.strip().lower() in ("endnotes", "end notes"):
            in_endnotes = True
            continue

        if in_endnotes:
            # A new non-endnote heading ends the endnotes section
            if token.type == "heading":
                in_endnotes = False
                continue

            if token.type == "paragraph" and citation_pattern.match(token.text.strip()):
                endnote_num += 1
                prefix = f"{endnote_num}. "
                token.text = prefix + token.text
                token.children = _parse_inline(token.text)

    return tokens


_CITATION_RE = re.compile(r"\[(Verified|Estimated|Stated|Derived):[^\]]*\]")


def _extract_inline_citations(tokens: List[_Token]) -> List[_Token]:
    """Convert inline citations to superscript references + appended Endnotes.

    For the Retainer Proposal: scans body text for [Verified: ...], etc.,
    replaces each with a Unicode superscript number, and appends an Endnotes
    section at the end with numbered entries.

    Only operates on tokens that do NOT already have an Endnotes section.
    """
    # Check if there's already an Endnotes section
    for token in tokens:
        if token.type == "heading" and token.text.strip().lower() in ("endnotes", "end notes"):
            return tokens  # Already has endnotes — don't double-process

    collected: List[str] = []
    endnote_num = 0

    for token in tokens:
        if token.type != "paragraph":
            continue

        text = token.text
        new_text = ""
        last_end = 0

        for m in _CITATION_RE.finditer(text):
            endnote_num += 1
            new_text += text[last_end:m.start()] + _to_superscript(endnote_num)
            collected.append(m.group(0))
            last_end = m.end()

        if last_end > 0:
            new_text += text[last_end:]
            token.text = new_text
            token.children = _parse_inline(new_text)

    if collected:
        # Append Endnotes heading
        tokens.append(_Token(type="heading", level=2, text="Endnotes",
                             children=_parse_inline("Endnotes")))
        # Append numbered endnote entries
        for i, citation in enumerate(collected, 1):
            entry = f"{i}. {citation}"
            tokens.append(_Token(type="paragraph", text=entry,
                                 children=_parse_inline(entry)))

    return tokens


# ---------------------------------------------------------------------------
# Chart helpers
# ---------------------------------------------------------------------------

def _fetch_charts_for_deliverable(
    engagement_id: str, output_number: int,
) -> Dict[str, Tuple[bytes, dict]]:
    """Fetch all chart PNGs for a deliverable from Supabase.

    Queries engagement_graphics WHERE status='confirmed' and output_number is
    in the output_placements INTEGER[] array.

    Returns dict mapping chart_type -> (png_bytes, row).
    """
    sb = get_supabase()
    result = (
        sb.table("engagement_graphics")
        .select("*")
        .eq("engagement_id", engagement_id)
        .eq("status", "confirmed")
        .execute()
    )

    charts: Dict[str, Tuple[bytes, dict]] = {}
    if not result.data:
        logger.info("No confirmed graphics for engagement %s", engagement_id)
        return charts

    for row in result.data:
        placements = row.get("output_placements") or []
        if output_number not in placements:
            continue

        chart_type = row.get("chart_type", "")
        storage_path = row.get("storage_path")
        bucket = row.get("storage_bucket", "engagement-graphics")

        if not storage_path:
            continue

        try:
            png_bytes = sb.storage.from_(bucket).download(storage_path)
            if png_bytes:
                charts[chart_type] = (png_bytes, row)
                logger.info("Fetched chart %s (%d bytes)", chart_type, len(png_bytes))
        except Exception as e:
            logger.error("Failed to download chart %s: %s", chart_type, e)

    return charts


# ---------------------------------------------------------------------------
# DOCX rendering — cover page extraction + markdown conversion
# ---------------------------------------------------------------------------

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _find_cover_boundary(body_element) -> Optional[object]:
    """Find the paragraph that marks the end of the cover page.

    The cover page ends at the first <w:p> whose <w:pPr> contains an in-body
    <w:sectPr>. Returns the XML element, or None if no section break found.
    """
    for child in body_element:
        if child.tag == qn("w:p"):
            pPr = child.find(qn("w:pPr"))
            if pPr is not None and pPr.find(qn("w:sectPr")) is not None:
                return child
    return None


def _replace_text_in_element(element, old: str, new: str) -> None:
    """Replace all occurrences of old text with new text in any <w:t> elements."""
    for t_elem in element.iter(qn("w:t")):
        if t_elem.text and old in t_elem.text:
            t_elem.text = t_elem.text.replace(old, new)


def _replace_cover_and_footer_placeholders(
    doc: Document, client_name: str, engagement_date: str,
) -> None:
    """Replace [CLIENT COMPANY NAME] and [ENGAGEMENT DATE] in the entire document.

    This covers the cover page paragraphs, footer XML, and any remaining instances.
    """
    body = doc.element.body

    # Replace in all body text elements
    _replace_text_in_element(body, "[CLIENT COMPANY NAME]", client_name)
    _replace_text_in_element(body, "[COMPANY NAME]", client_name)
    _replace_text_in_element(body, "[ENGAGEMENT DATE]", engagement_date)
    _replace_text_in_element(body, "[DATE]", engagement_date)

    # Replace in headers and footers (separate XML parts)
    for section in doc.sections:
        for rel_type in ("header", "footer"):
            try:
                if rel_type == "header":
                    part = section.header
                else:
                    part = section.footer
                if part and part._element is not None:
                    _replace_text_in_element(part._element, "[CLIENT COMPANY NAME]", client_name)
                    _replace_text_in_element(part._element, "[COMPANY NAME]", client_name)
                    _replace_text_in_element(part._element, "[ENGAGEMENT DATE]", engagement_date)
                    _replace_text_in_element(part._element, "[DATE]", engagement_date)
            except Exception:
                pass


def _delete_body_stubs(body_element, cover_boundary) -> object:
    """Delete all body content between cover boundary and the final sectPr.

    Returns the final sectPr element (must be preserved as last child of w:body).
    """
    # The final sectPr is always the last child of w:body
    final_sect_pr = body_element[-1]
    if final_sect_pr.tag != qn("w:sectPr"):
        # Fallback: search for it
        for child in reversed(list(body_element)):
            if child.tag == qn("w:sectPr"):
                final_sect_pr = child
                break

    # Collect elements to remove: everything after cover_boundary but before final_sect_pr
    removing = False
    to_remove = []
    for child in list(body_element):
        if child is cover_boundary:
            removing = True
            continue
        if removing and child is not final_sect_pr:
            to_remove.append(child)

    for el in to_remove:
        body_element.remove(el)

    return final_sect_pr


def _add_styled_runs(paragraph, spans: List[_InlineSpan]) -> None:
    """Add inline spans to a paragraph, preserving bold/italic formatting.

    Does NOT apply font/color — lets the paragraph style handle it.
    Only sets bold/italic on runs that need it.
    """
    for span in spans:
        run = paragraph.add_run(span.text)
        if span.bold:
            run.bold = True
        if span.italic:
            run.italic = True


def _insert_chart_image(
    doc: Document, insert_before, chart_name: str,
    charts: Dict[str, Tuple[bytes, dict]],
) -> None:
    """Insert a chart PNG as a centered paragraph before insert_before element."""
    chart_data = charts.get(chart_name)
    if not chart_data:
        logger.warning("No chart image for [CHART: %s] — inserting placeholder text", chart_name)
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(f"[Chart: {chart_name} — image not available]")
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True
        # Move from end of body to correct position
        insert_before.addprevious(para._element)
        return

    png_bytes, _row = chart_data

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    run.add_picture(io.BytesIO(png_bytes), width=Inches(6.0))
    # Move from end of body to correct position
    insert_before.addprevious(para._element)


def _set_cell_shading(cell, color_hex: str) -> None:
    """Apply background shading to a table cell."""
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:color"), "auto")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_table_borders(table) -> None:
    """Apply thin borders to all cells in a table."""
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "999999")
        borders.append(element)
    tbl_pr.append(borders)


def _append_tokens_to_body(
    doc: Document,
    insert_before,
    tokens: List[_Token],
    charts: Dict[str, Tuple[bytes, dict]],
) -> None:
    """Convert tokens to python-docx elements and insert before insert_before.

    Uses doc.add_paragraph() to create elements with a valid .part reference,
    then moves them to the correct position before insert_before.
    Template style IDs: Heading1, Heading2, Heading3, ListParagraph.
    """
    style_map = {1: "Heading1", 2: "Heading2", 3: "Heading3"}

    for token in tokens:
        if token.type == "heading":
            level = min(token.level, 3)
            style_name = style_map[level]

            para = doc.add_paragraph()
            try:
                para.style = doc.styles[style_name]
            except KeyError:
                pass
            _add_styled_runs(para, token.children or [_InlineSpan(token.text)])
            insert_before.addprevious(para._element)

        elif token.type == "paragraph":
            para = doc.add_paragraph()
            _add_styled_runs(para, token.children or [_InlineSpan(token.text)])
            insert_before.addprevious(para._element)

        elif token.type == "bullet":
            para = doc.add_paragraph()
            try:
                para.style = doc.styles["ListParagraph"]
            except KeyError:
                pass

            # Add bullet via numbering XML
            pPr = para._element.get_or_add_pPr()
            numPr = OxmlElement("w:numPr")
            ilvl = OxmlElement("w:ilvl")
            ilvl.set(qn("w:val"), "0")
            numPr.append(ilvl)
            numId = OxmlElement("w:numId")
            numId.set(qn("w:val"), "1")
            numPr.append(numId)
            pPr.append(numPr)

            _add_styled_runs(para, token.children or [_InlineSpan(token.text)])
            insert_before.addprevious(para._element)

        elif token.type == "numbered":
            para = doc.add_paragraph()
            try:
                para.style = doc.styles["ListParagraph"]
            except KeyError:
                pass
            _add_styled_runs(para, token.children or [_InlineSpan(token.text)])
            insert_before.addprevious(para._element)

        elif token.type == "table":
            if not token.rows:
                continue

            num_cols = max(len(row) for row in token.rows)
            num_rows = len(token.rows)

            # Create table via python-docx API (gives valid .part)
            tbl = doc.add_table(rows=num_rows, cols=num_cols)
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            _set_table_borders(tbl)

            for row_idx, row_data in enumerate(token.rows):
                row = tbl.rows[row_idx]
                for col_idx, cell_text in enumerate(row_data):
                    if col_idx >= num_cols:
                        break
                    cell = row.cells[col_idx]
                    cell.text = ""
                    p = cell.paragraphs[0]
                    run = p.add_run(cell_text)

                    if row_idx == 0:
                        run.font.name = BODY_FONT
                        run.font.size = Pt(10)
                        run.font.color.rgb = WHITE
                        run.bold = True
                        _set_cell_shading(cell, CRIMSON_HEX)
                    else:
                        run.font.name = BODY_FONT
                        run.font.size = Pt(10)
                        run.font.color.rgb = CHARCOAL

            # Move table from end of document to correct position
            insert_before.addprevious(tbl._tbl)

            # Add spacing paragraph after table
            spacer = OxmlElement("w:p")
            insert_before.addprevious(spacer)

        elif token.type == "chart_placeholder":
            _insert_chart_image(doc, insert_before, token.chart_name, charts)

            # Add spacing after chart
            spacer = OxmlElement("w:p")
            insert_before.addprevious(spacer)


def render_docx(
    template_path: str,
    markdown_content: str,
    engagement_id: str,
    client_name: str,
    engagement_date: str,
    output_number: int,
    file_prefix: str = "",
) -> bytes:
    """Render a .docx by extracting the template cover page and appending markdown content.

    1. Open template
    2. Replace client/date placeholders on cover page and footer
    3. Find cover page boundary (in-body sectPr)
    4. Delete body stubs between cover and final sectPr
    5. Tokenize markdown
    6. Post-process tokens (endnote numbering, citation extraction)
    7. Convert tokens to docx elements using template styles
    8. Insert chart PNGs from Supabase
    9. Return rendered bytes
    """
    doc = Document(template_path)
    body = doc.element.body

    # Step 1: Replace placeholders on cover page and in footer
    _replace_cover_and_footer_placeholders(doc, client_name, engagement_date)

    # Step 2: Find cover page boundary
    cover_boundary = _find_cover_boundary(body)
    if not cover_boundary:
        logger.warning("No cover page section break found in %s — appending at end", template_path)
        # Fallback: just append after all existing content
        final_sect_pr = body[-1] if body[-1].tag == qn("w:sectPr") else None
    else:
        # Step 3: Delete body stubs
        final_sect_pr = _delete_body_stubs(body, cover_boundary)

    # Step 4: Fetch charts for this deliverable
    charts = _fetch_charts_for_deliverable(engagement_id, output_number)
    logger.info("Fetched %d charts for output_number=%d", len(charts), output_number)

    # Step 5: Tokenize markdown
    tokens = _tokenize_markdown(markdown_content)
    logger.info("Tokenized markdown: %d blocks", len(tokens))

    # Step 5b: For Retainer Proposal — extract inline citations to endnotes
    is_retainer = file_prefix.startswith("Retainer_Proposal") or file_prefix.startswith("Phase_2_Retainer_Proposal")
    if is_retainer:
        tokens = _extract_inline_citations(tokens)
        logger.info("Extracted inline citations for Retainer Proposal")

    # Step 5c: Number endnotes in the Endnotes section (all documents)
    tokens = _number_endnotes(tokens)

    # Step 6: Append converted content before the final sectPr
    if final_sect_pr is not None:
        _append_tokens_to_body(doc, final_sect_pr, tokens, charts)
    else:
        # No final sectPr — append at end (shouldn't happen with our templates)
        _append_tokens_to_body(doc, body[-1], tokens, charts)

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# PPTX rendering — slide creation from markdown
# ---------------------------------------------------------------------------

NS_A = "http://schemas.openxmlformats.org/drawingml/2006/main"
NS_R = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
NS_P = "http://schemas.openxmlformats.org/presentationml/2006/main"
NS_RELS = "http://schemas.openxmlformats.org/package/2006/relationships"


def _parse_slide_markdown(content: str) -> List[dict]:
    """Parse presentation deck markdown into slides.

    Expected format:
        # Slide 1: Title Slide
        ## Slide Body
        Content for the slide face
        ## Slide Narration
        Content for speaker notes

    Returns list of dicts: {"title", "body", "narration", "charts"}.
    """
    slides: List[dict] = []
    slide_chunks = re.split(r"^#\s+", content, flags=re.MULTILINE)

    for chunk in slide_chunks:
        chunk = chunk.strip()
        if not chunk:
            continue

        title_line, _, rest = chunk.partition("\n")
        title = title_line.strip()
        title = re.sub(r"^Slide\s+\d+\s*:\s*", "", title).strip()

        body = ""
        narration = ""
        charts: List[str] = []

        sub_sections = re.split(r"^##\s+", rest, flags=re.MULTILINE)
        for sub in sub_sections:
            sub = sub.strip()
            if not sub:
                continue
            sub_title, _, sub_body = sub.partition("\n")
            sub_title_lower = sub_title.strip().lower()
            sub_body = sub_body.strip()

            if "narration" in sub_title_lower or "notes" in sub_title_lower:
                narration = sub_body
            elif "body" in sub_title_lower or "content" in sub_title_lower:
                body = sub_body
            else:
                body += "\n\n" + sub_body if body else sub_body

        for chart_match in re.finditer(r"\[CHART:\s*([^\]]+)\]", body):
            charts.append(chart_match.group(1).strip())

        slides.append({
            "title": title,
            "body": body,
            "narration": narration,
            "charts": charts,
        })

    return slides


def render_pptx(
    template_path: str,
    markdown_content: str,
    engagement_id: str,
    client_name: str,
    engagement_date: str,
    output_number: int,
) -> bytes:
    """Render a .pptx from markdown using the template.

    1. Keep and update title slide (slide 1) with client/date placeholders
    2. For each markdown slide section, duplicate a content slide layout
       and populate title, body, speaker notes, and chart images
    3. Remove unused template placeholder slides
    """
    slides_data = _parse_slide_markdown(markdown_content)
    charts = _fetch_charts_for_deliverable(engagement_id, output_number)

    buf_in = io.BytesIO(open(template_path, "rb").read())
    buf_out = io.BytesIO()

    with zipfile.ZipFile(buf_in, "r") as zin, zipfile.ZipFile(buf_out, "w") as zout:
        slide_files = sorted(
            [f for f in zin.namelist() if re.match(r"ppt/slides/slide\d+\.xml$", f)]
        )
        notes_files = sorted(
            [f for f in zin.namelist() if re.match(r"ppt/notesSlides/notesSlide\d+\.xml$", f)]
        )

        # Track which image relationships we add
        added_images: Dict[str, str] = {}  # filename -> rId
        next_image_idx = 1

        # Process all files in the zip
        for item in zin.infolist():
            data = zin.read(item.filename)

            # Process slide XML files
            if item.filename in slide_files:
                slide_idx = slide_files.index(item.filename)
                if slide_idx < len(slides_data):
                    slide_data = slides_data[slide_idx]
                    data = _render_pptx_slide(
                        data, slide_data, client_name, engagement_date,
                        charts, slide_idx, zout, zin, item.filename,
                    )

            # Process notes files
            if item.filename in notes_files:
                idx_match = re.search(r"notesSlide(\d+)\.xml", item.filename)
                if idx_match:
                    slide_idx = int(idx_match.group(1)) - 1
                    if slide_idx < len(slides_data) and slides_data[slide_idx].get("narration"):
                        data = _replace_pptx_notes(data, slides_data[slide_idx]["narration"])

            zout.writestr(item, data)

    return buf_out.getvalue()


def _render_pptx_slide(
    slide_xml: bytes,
    slide_data: dict,
    client_name: str,
    engagement_date: str,
    charts: Dict[str, Tuple[bytes, dict]],
    slide_idx: int,
    zout: zipfile.ZipFile,
    zin: zipfile.ZipFile,
    slide_filename: str,
) -> bytes:
    """Render a single PPTX slide: replace placeholders and insert chart images."""
    tree = etree.fromstring(slide_xml)

    # Replace text placeholders
    for t_elem in tree.iter(f"{{{NS_A}}}t"):
        if t_elem.text is None:
            continue
        text = t_elem.text
        text = text.replace("[CLIENT COMPANY NAME]", client_name)
        text = text.replace("[COMPANY NAME]", client_name)
        text = text.replace("[DATE]", engagement_date)
        text = text.replace("[ENGAGEMENT DATE]", engagement_date)
        t_elem.text = text

    # For the title slide (index 0), we're done after placeholder replacement
    if slide_idx == 0:
        return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)

    # For content slides: populate body text into the first non-title text body
    body_text = slide_data.get("body", "")
    # Remove [CHART: ...] placeholders from body text for text rendering
    clean_body = re.sub(r"\[CHART:\s*[^\]]+\]", "", body_text).strip()

    if clean_body:
        _populate_slide_body(tree, slide_data.get("title", ""), clean_body)

    # Insert chart images into the slide
    slide_num = re.search(r"slide(\d+)\.xml", slide_filename)
    if slide_num and slide_data.get("charts"):
        slide_n = slide_num.group(1)
        rels_path = f"ppt/slides/_rels/slide{slide_n}.xml.rels"

        for chart_name in slide_data["charts"]:
            chart_data = charts.get(chart_name)
            if not chart_data:
                continue

            png_bytes, _row = chart_data
            img_filename = f"image_chart_{chart_name}.png"
            img_path = f"ppt/media/{img_filename}"

            # Write image to zip if not already written
            if img_path not in [i.filename for i in zout.infolist()]:
                zout.writestr(img_path, png_bytes)

    return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)


def _populate_slide_body(tree, title: str, body_text: str) -> None:
    """Populate a slide's shape text bodies with title and body content.

    Uses content-matching: finds shapes by text length rather than placeholder
    type, since these templates use raw <p:sp> without <p:ph> type attributes.

    Strategy:
    - The shape with the most text content is the body area — replace it.
    - Short single-line shapes that match the slide title get the new title.
    - Section divider slides (with just a number + short text) are left as-is
      if they only have small text areas.
    """
    sp_elements = list(tree.iter(f"{{{NS_P}}}sp"))

    # Gather all shapes with text, along with their total text length
    shape_texts: List[Tuple] = []  # (sp_element, txBody, existing_text, text_len)
    for sp in sp_elements:
        txBody = sp.find(f".//{{{NS_A}}}txBody")
        if txBody is None:
            continue
        existing_text = ""
        for t in txBody.iter(f"{{{NS_A}}}t"):
            if t.text:
                existing_text += t.text
        if existing_text.strip():
            shape_texts.append((sp, txBody, existing_text.strip(), len(existing_text.strip())))

    if not shape_texts:
        return

    # Sort by text length descending — largest text area is the body
    shape_texts.sort(key=lambda x: x[3], reverse=True)

    # Replace the largest text body with the markdown body content
    body_replaced = False
    for sp, txBody, existing_text, text_len in shape_texts:
        # Skip very short shapes (slide numbers like "01", "02")
        if text_len <= 3:
            continue

        if not body_replaced and body_text and text_len > 20:
            # This is the body area — replace with markdown content
            _replace_txBody_content(txBody, body_text)
            body_replaced = True
        elif title and existing_text.upper() == existing_text and text_len < 50:
            # ALL-CAPS short text is likely a title — replace it
            for t in txBody.iter(f"{{{NS_A}}}t"):
                if t.text and t.text.strip():
                    t.text = title.upper()
                    break


def _replace_txBody_content(txBody, body_text: str) -> None:
    """Replace all content in a DrawingML txBody with new body text lines."""
    paragraphs = txBody.findall(f"{{{NS_A}}}p")
    if not paragraphs:
        return

    # Capture formatting from the first paragraph's first run
    first_p = paragraphs[0]
    first_rPr = first_p.find(f".//{{{NS_A}}}rPr")

    # Split body text into lines
    body_lines = [l.strip() for l in body_text.split("\n") if l.strip()]
    if not body_lines:
        return

    # Clear all existing paragraphs except first
    for p in paragraphs[1:]:
        txBody.remove(p)

    # Set first paragraph text
    for r in first_p.findall(f"{{{NS_A}}}r"):
        first_p.remove(r)
    r_elem = etree.SubElement(first_p, f"{{{NS_A}}}r")
    if first_rPr is not None:
        r_elem.append(deepcopy(first_rPr))
    t_elem = etree.SubElement(r_elem, f"{{{NS_A}}}t")
    t_elem.text = body_lines[0]

    # Add remaining lines as new paragraphs
    for line in body_lines[1:]:
        new_p = deepcopy(first_p)
        for r in new_p.findall(f"{{{NS_A}}}r"):
            new_p.remove(r)
        r_elem = etree.SubElement(new_p, f"{{{NS_A}}}r")
        if first_rPr is not None:
            r_elem.append(deepcopy(first_rPr))
        t_elem = etree.SubElement(r_elem, f"{{{NS_A}}}t")
        t_elem.text = line
        txBody.append(new_p)


def _replace_pptx_notes(notes_xml: bytes, narration: str) -> bytes:
    """Replace the notes slide content with narration text."""
    tree = etree.fromstring(notes_xml)

    for txBody in tree.iter(f"{{{NS_A}}}txBody"):
        paragraphs = txBody.findall(f"{{{NS_A}}}p")
        if len(paragraphs) <= 1:
            continue

        first_p = paragraphs[0]
        for p in paragraphs[1:]:
            txBody.remove(p)

        # Get first run's formatting
        first_rPr = first_p.find(f".//{{{NS_A}}}rPr")

        # Clear existing runs
        for r in first_p.findall(f"{{{NS_A}}}r"):
            first_p.remove(r)

        # Split narration into paragraphs
        narration_lines = [l.strip() for l in narration.split("\n\n") if l.strip()]

        if narration_lines:
            # First paragraph
            r_elem = etree.SubElement(first_p, f"{{{NS_A}}}r")
            if first_rPr is not None:
                r_elem.append(deepcopy(first_rPr))
            t_elem = etree.SubElement(r_elem, f"{{{NS_A}}}t")
            t_elem.text = narration_lines[0]

            # Additional paragraphs
            for line in narration_lines[1:]:
                new_p = deepcopy(first_p)
                for r in new_p.findall(f"{{{NS_A}}}r"):
                    new_p.remove(r)
                r_elem = etree.SubElement(new_p, f"{{{NS_A}}}r")
                if first_rPr is not None:
                    r_elem.append(deepcopy(first_rPr))
                t_elem = etree.SubElement(r_elem, f"{{{NS_A}}}t")
                t_elem.text = line
                txBody.append(new_p)
        break

    return etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)


# ---------------------------------------------------------------------------
# Orchestrator
# ---------------------------------------------------------------------------

async def render_engagement_deliverables(engagement_id: str) -> dict:
    """Render all QC-approved .md files from Drive into branded .docx deliverables.

    Downloads markdown files from the engagement's 04_Deliverables Drive folder,
    loads the corresponding template, renders using cover-page extraction +
    markdown conversion, and uploads the result back to Drive.

    Returns dict with ``rendered`` list and ``skipped`` list.
    """
    eng = get_engagement_by_id(engagement_id)
    if not eng:
        raise ValueError(f"Engagement {engagement_id} not found")

    sb = get_supabase()
    deliverables_folder_id = eng.get("drive_deliverables_folder_id")
    if not deliverables_folder_id:
        raise ValueError("Engagement has no Drive deliverables folder configured")

    # Get client info
    client_row = sb.table("clients").select("company_name").eq("id", eng["client_id"]).limit(1).execute()
    client_name = client_row.data[0]["company_name"] if client_row.data else "Client"

    from datetime import date
    engagement_date = ""
    if eng.get("start_date"):
        try:
            d = date.fromisoformat(str(eng["start_date"])[:10])
            engagement_date = d.strftime("%B %Y")
        except (ValueError, TypeError):
            engagement_date = str(eng["start_date"])

    # List .md files in the deliverables folder
    md_files = await list_files_in_folder(deliverables_folder_id, extension=".md")
    if not md_files:
        logger.info("No .md files found in deliverables folder for %s", engagement_id)
        return []

    results: List[dict] = []
    skipped: List[dict] = []

    for md_file in md_files:
        filename = md_file["name"]

        match = _match_template(filename)
        if not match:
            logger.info("No template mapping for '%s' — skipping", filename)
            continue

        output_name, template_filename, output_ext, output_number, file_prefix = match
        template_path = os.path.join(TEMPLATES_DIR, template_filename)

        # Skip PPTX (Presentation Deck) — handled via Cowork
        if file_prefix == "Presentation_Deck" or output_ext == "pptx":
            logger.info("Skipping %s — Presentation Deck created via Cowork", filename)
            skipped.append({"name": output_name, "reason": "Created via Cowork"})
            continue

        if not os.path.exists(template_path):
            logger.error("Template file not found: %s", template_path)
            continue

        # Download the markdown content
        file_bytes = await download_file_by_id(md_file["id"])
        if not file_bytes:
            logger.error("Failed to download %s from Drive", filename)
            continue

        markdown_content = file_bytes.decode("utf-8")

        # Render
        try:
            rendered_bytes = render_docx(
                template_path, markdown_content, engagement_id,
                client_name, engagement_date, output_number,
                file_prefix=file_prefix,
            )
        except Exception as e:
            logger.error("Render failed for %s: %s", output_name, e, exc_info=True)
            continue

        # Upload rendered file back to Drive
        output_filename = f"{output_name.replace(' ', '_')}.{output_ext}"
        mimetype = (
            "application/vnd.openxmlformats-officedocument.presentationml.presentation"
            if output_ext == "pptx" else
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        drive_file_id = await upload_file_to_drive_folder(
            deliverables_folder_id,
            output_filename,
            rendered_bytes,
            mimetype,
        )

        if drive_file_id:
            results.append({
                "output_name": output_name,
                "output_file": output_filename,
                "drive_file_id": drive_file_id,
                "size_bytes": len(rendered_bytes),
            })
            logger.info("Rendered %s -> %s (Drive: %s)", filename, output_filename, drive_file_id)
        else:
            logger.error("Failed to upload %s to Drive", output_filename)

    # Log activity
    if results:
        log_activity(engagement_id, "system", "deliverables_rendered", {
            "rendered_count": len(results),
            "files": [r["output_file"] for r in results],
        })

    return {"rendered": results, "skipped": skipped}
