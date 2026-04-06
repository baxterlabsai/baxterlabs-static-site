"""Tests for the Phase 7 renderer inline markdown normalizer and verification gate.

Added 2026-04-06 after the Scion Phase 7 endnote leakage incident.
"""
from __future__ import annotations

import io
import re
import zipfile

import pytest
from docx import Document
from docx.shared import Pt

from services.document_renderer import (
    _parse_inline,
    _add_styled_runs,
    _verify_docx_no_residual_markdown,
    _verify_pptx_no_residual_markdown,
    RendererVerificationError,
)


# ---------------------------------------------------------------------------
# Normalizer tests
# ---------------------------------------------------------------------------

class TestParseInline:
    """Tests for _parse_inline superscript handling."""

    def test_pandoc_superscript_basic(self):
        spans = _parse_inline("Text with ^1^ marker.")
        types = [(s.text, s.superscript) for s in spans]
        assert ("1", True) in types
        assert all(not s.superscript for s in spans if s.text != "1")

    def test_markdown_footnote_basic(self):
        spans = _parse_inline("Text with [^42] marker.")
        types = [(s.text, s.superscript) for s in spans]
        assert ("42", True) in types

    def test_both_dialects_in_one_string(self):
        spans = _parse_inline("First ^1^ and second [^2] here.")
        supers = [s for s in spans if s.superscript]
        assert len(supers) == 2
        assert supers[0].text == "1"
        assert supers[1].text == "2"

    def test_three_digit_number(self):
        spans = _parse_inline("Reference ^160^ in text.")
        supers = [s for s in spans if s.superscript]
        assert len(supers) == 1
        assert supers[0].text == "160"

    def test_marker_at_start(self):
        spans = _parse_inline("^1^ at start.")
        assert spans[0].superscript is True
        assert spans[0].text == "1"

    def test_marker_at_end(self):
        spans = _parse_inline("End of sentence.^5^")
        supers = [s for s in spans if s.superscript]
        assert len(supers) == 1
        assert supers[0].text == "5"

    def test_multiple_markers(self):
        spans = _parse_inline("A^1^ B^2^ C[^3] D.")
        supers = [s for s in spans if s.superscript]
        assert len(supers) == 3

    def test_bold_and_superscript_mixed(self):
        spans = _parse_inline("**Bold** and ^1^ ref.")
        bold_spans = [s for s in spans if s.bold]
        super_spans = [s for s in spans if s.superscript]
        assert len(bold_spans) == 1
        assert bold_spans[0].text == "Bold"
        assert len(super_spans) == 1
        assert super_spans[0].text == "1"

    def test_plain_text_unchanged(self):
        spans = _parse_inline("No markers here at all.")
        assert len(spans) == 1
        assert spans[0].text == "No markers here at all."
        assert not spans[0].superscript
        assert not spans[0].bold
        assert not spans[0].italic

    def test_surrounding_text_preserved(self):
        spans = _parse_inline("before ^7^ after")
        texts = [s.text for s in spans]
        assert "before " in texts
        assert "7" in texts
        assert " after" in texts


class TestAddStyledRuns:
    """Tests for _add_styled_runs with superscript."""

    def test_superscript_run_created(self):
        doc = Document()
        para = doc.add_paragraph()
        spans = _parse_inline("Text with ^1^ and [^2] markers.")
        _add_styled_runs(para, spans)

        super_runs = [r for r in para.runs if r.font.superscript]
        assert len(super_runs) == 2
        assert super_runs[0].text == "1"
        assert super_runs[1].text == "2"

    def test_no_literal_markers_in_output(self):
        doc = Document()
        para = doc.add_paragraph()
        spans = _parse_inline("End note.^42^ Another.[^7]")
        _add_styled_runs(para, spans)

        full_text = para.text
        assert "^42^" not in full_text
        assert "[^7]" not in full_text
        assert "42" in full_text
        assert "7" in full_text

    def test_superscript_ooxml_vertAlign(self):
        """Verify the actual OOXML contains vertAlign superscript."""
        doc = Document()
        para = doc.add_paragraph()
        spans = _parse_inline("Sample ^1^ text.")
        _add_styled_runs(para, spans)

        xml = para._element.xml
        assert "w:vertAlign" in xml
        assert 'w:val="superscript"' in xml

    def test_table_cell_font_params(self):
        """Verify font_name/font_size/font_color params work."""
        from docx.shared import RGBColor
        doc = Document()
        para = doc.add_paragraph()
        spans = _parse_inline("Cell ^1^ text")
        _add_styled_runs(para, spans, font_name="Arial", font_size=10,
                         font_color=RGBColor(0x2D, 0x34, 0x36))

        for run in para.runs:
            assert run.font.name == "Arial"
            assert run.font.size == Pt(10)

    def test_force_bold_header_row(self):
        """Verify force_bold makes all runs bold (table header)."""
        doc = Document()
        para = doc.add_paragraph()
        spans = _parse_inline("Header ^1^ text")
        _add_styled_runs(para, spans, force_bold=True)

        for run in para.runs:
            assert run.bold is True


# ---------------------------------------------------------------------------
# Verification gate tests — DOCX
# ---------------------------------------------------------------------------

def _make_docx_with_text(text: str) -> bytes:
    """Helper: create a minimal DOCX with one paragraph containing the given text."""
    doc = Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestVerifyDocxGate:
    """Tests for _verify_docx_no_residual_markdown."""

    def test_clean_docx_passes(self):
        """No markdown → no error."""
        docx_bytes = _make_docx_with_text("This is clean text with no markdown.")
        _verify_docx_no_residual_markdown(docx_bytes, "test_clean")

    def test_pandoc_superscript_detected(self):
        docx_bytes = _make_docx_with_text("Text with ^1^ leaked marker.")
        with pytest.raises(RendererVerificationError, match="pandoc_superscript"):
            _verify_docx_no_residual_markdown(docx_bytes, "test_pandoc")

    def test_markdown_footnote_detected(self):
        docx_bytes = _make_docx_with_text("Text with [^42] leaked marker.")
        with pytest.raises(RendererVerificationError, match="markdown_footnote"):
            _verify_docx_no_residual_markdown(docx_bytes, "test_footnote")

    def test_bold_markdown_detected(self):
        docx_bytes = _make_docx_with_text("This has **bold text** that leaked.")
        with pytest.raises(RendererVerificationError, match="markdown_bold"):
            _verify_docx_no_residual_markdown(docx_bytes, "test_bold")

    def test_italic_markdown_detected(self):
        docx_bytes = _make_docx_with_text("This has *italic text* that leaked.")
        with pytest.raises(RendererVerificationError, match="markdown_italic"):
            _verify_docx_no_residual_markdown(docx_bytes, "test_italic")

    def test_code_span_detected(self):
        docx_bytes = _make_docx_with_text("This has `code span` that leaked.")
        with pytest.raises(RendererVerificationError, match="code_span"):
            _verify_docx_no_residual_markdown(docx_bytes, "test_code")

    def test_markdown_link_detected(self):
        docx_bytes = _make_docx_with_text("This has [link](http://example.com) that leaked.")
        with pytest.raises(RendererVerificationError, match="markdown_link"):
            _verify_docx_no_residual_markdown(docx_bytes, "test_link")

    def test_normalizer_then_gate_passes(self):
        """Text that goes through the normalizer should pass the gate."""
        doc = Document()
        para = doc.add_paragraph()
        spans = _parse_inline("Test paragraph with ^1^ marker and [^2] too.")
        _add_styled_runs(para, spans)

        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        # Should NOT raise — the normalizer converted markers to superscript runs
        _verify_docx_no_residual_markdown(docx_bytes, "test_normalized")

    def test_table_cell_leakage_detected(self):
        """Markdown in a table cell should also be caught."""
        doc = Document()
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "Cell with ^3^ marker."
        buf = io.BytesIO()
        doc.save(buf)
        docx_bytes = buf.getvalue()

        with pytest.raises(RendererVerificationError, match="pandoc_superscript"):
            _verify_docx_no_residual_markdown(docx_bytes, "test_table_cell")

    def test_error_message_contains_context(self):
        """Error message should include pattern name, count, and context."""
        docx_bytes = _make_docx_with_text("First ^1^ and second ^2^ markers.")
        with pytest.raises(RendererVerificationError) as exc_info:
            _verify_docx_no_residual_markdown(docx_bytes, "test_context")
        msg = str(exc_info.value)
        assert "pandoc_superscript" in msg
        assert "occurrence" in msg
        assert "test_context" in msg
        # Count may be >2 because both python-docx walk and raw XML fallback
        # find the same text — this is intentionally defensive (catches text boxes)


# ---------------------------------------------------------------------------
# Verification gate tests — PPTX
# ---------------------------------------------------------------------------

def _make_pptx_with_text(text: str) -> bytes:
    """Helper: create a minimal PPTX with one slide containing the given text."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w") as z:
        # Minimal [Content_Types].xml
        z.writestr("[Content_Types].xml",
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
            '<Default Extension="xml" ContentType="application/xml"/>'
            '</Types>')
        # Minimal slide with text
        slide_xml = (
            '<?xml version="1.0" encoding="UTF-8"?>'
            '<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"'
            '       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
            '<p:cSld><p:spTree><p:sp><p:txBody><a:p><a:r>'
            f'<a:t>{text}</a:t>'
            '</a:r></a:p></p:txBody></p:sp></p:spTree></p:cSld></p:sld>'
        )
        z.writestr("ppt/slides/slide1.xml", slide_xml)
    return buf.getvalue()


class TestVerifyPptxGate:
    """Tests for _verify_pptx_no_residual_markdown."""

    def test_clean_pptx_passes(self):
        pptx_bytes = _make_pptx_with_text("Clean text with no markdown.")
        _verify_pptx_no_residual_markdown(pptx_bytes, "test_clean_pptx")

    def test_pandoc_superscript_detected(self):
        pptx_bytes = _make_pptx_with_text("Slide text with ^1^ marker.")
        with pytest.raises(RendererVerificationError, match="pandoc_superscript"):
            _verify_pptx_no_residual_markdown(pptx_bytes, "test_pptx_pandoc")

    def test_markdown_footnote_detected(self):
        pptx_bytes = _make_pptx_with_text("Slide text with [^42] marker.")
        with pytest.raises(RendererVerificationError, match="markdown_footnote"):
            _verify_pptx_no_residual_markdown(pptx_bytes, "test_pptx_footnote")

    def test_bold_markdown_detected(self):
        pptx_bytes = _make_pptx_with_text("Slide with **bold** leak.")
        with pytest.raises(RendererVerificationError, match="markdown_bold"):
            _verify_pptx_no_residual_markdown(pptx_bytes, "test_pptx_bold")

    def test_italic_markdown_detected(self):
        pptx_bytes = _make_pptx_with_text("Slide with *italic* leak.")
        with pytest.raises(RendererVerificationError, match="markdown_italic"):
            _verify_pptx_no_residual_markdown(pptx_bytes, "test_pptx_italic")

    def test_notes_slide_checked(self):
        """Markdown in notes slides should also be caught."""
        buf = io.BytesIO()
        with zipfile.ZipFile(buf, "w") as z:
            z.writestr("[Content_Types].xml",
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types">'
                '<Default Extension="xml" ContentType="application/xml"/>'
                '</Types>')
            z.writestr("ppt/slides/slide1.xml",
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<p:sld xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"'
                '       xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
                '<p:cSld><p:spTree/></p:cSld></p:sld>')
            z.writestr("ppt/notesSlides/notesSlide1.xml",
                '<?xml version="1.0" encoding="UTF-8"?>'
                '<p:notes xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"'
                '         xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">'
                '<p:cSld><p:spTree><p:sp><p:txBody><a:p><a:r>'
                '<a:t>Note with ^5^ marker.</a:t>'
                '</a:r></a:p></p:txBody></p:sp></p:spTree></p:cSld></p:notes>')
        pptx_bytes = buf.getvalue()

        with pytest.raises(RendererVerificationError, match="pandoc_superscript"):
            _verify_pptx_no_residual_markdown(pptx_bytes, "test_pptx_notes")
